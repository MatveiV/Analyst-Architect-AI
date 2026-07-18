"""
Export Service — экспорт данных в JSON, CSV, DOCX, PDF.
"""
import csv
import io
import json
import os
from datetime import datetime
from typing import Any

# Эпик B6: человекочитаемые названия стандартов для титульного блока экспорта
STANDARD_LABELS = {
    "IEEE_830": "IEEE 830-1998 (SRS)",
    "ISO_IEC_IEEE_29148": "ISO/IEC/IEEE 29148 (Requirements Engineering)",
    "GOST_19": "ГОСТ 19.201-78 (ТЗ на программу, ЕСПД)",
    "GOST_34": "ГОСТ 34.602-2020 (ТЗ на автоматизированную систему)",
    "C4_MODEL": "C4-модель (Simon Brown)",
    "UML_ISO_19505": "UML по ISO/IEC 19505",
    "ISO_IEC_IEEE_42010": "ISO/IEC/IEEE 42010 (Architecture description)",
    "GOST_19_701": "ГОСТ 19.701-90 (Схемы алгоритмов, программ, данных)",
    "IEC_61082": "IEC 61082 (Оформление технической документации)",
}

DIAGRAM_TYPE_LABELS = {
    "c4_context": "C4 — Контекст (Context)",
    "c4_container": "C4 — Контейнеры (Container)",
    "c4_component": "C4 — Компоненты (Component)",
    "use_case": "UML — Use Case",
    "sequence": "UML — Sequence",
    "class": "UML — Class",
    "erd": "ER-диаграмма (данные)",
    "flowchart": "Блок-схема процесса",
}


def export_review_json(review_data: dict) -> bytes:
    return json.dumps(review_data, ensure_ascii=False, indent=2).encode("utf-8")


def export_review_csv(review_data: dict) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow(["Поле", "Значение"])
    writer.writerow(["summary", review_data.get("summary", "")])
    writer.writerow(["confidence", review_data.get("confidence", "")])
    writer.writerow(["needs_review", review_data.get("needs_review", "")])

    writer.writerow([])
    writer.writerow(["Риски", ""])
    writer.writerow(["severity", "description"])
    for risk in review_data.get("risks", []):
        writer.writerow([risk.get("severity", ""), risk.get("description", "")])

    writer.writerow([])
    writer.writerow(["Вопросы заказчику", ""])
    for q in review_data.get("questions_to_client", []):
        writer.writerow([q])

    writer.writerow([])
    writer.writerow(["Критерии приёмки", ""])
    for c in review_data.get("acceptance_criteria", []):
        writer.writerow([c])

    return output.getvalue().encode("utf-8-sig")  # BOM for Excel


def export_document_docx(title: str, content: dict, diagrams: list | None = None) -> bytes:
    """
    Export document to DOCX format.

    Эпик A4: если передан список diagrams (объекты DiagramArtifact с render_png),
    диаграммы встраиваются в документ как изображения, а не только текстом кода.
    Эпик B6: если в content есть standard_profile — на титульном листе явно указывается
    применённый стандарт оформления.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, level=1)

        # Эпик B6: явное указание применённого стандарта
        standard = content.get("standard_profile")
        if standard:
            label = STANDARD_LABELS.get(standard, standard)
            note = doc.add_paragraph()
            note.add_run(f"Документ сформирован в соответствии со стандартом: {label}").italic = True

        def add_section(heading: str, items: list, bullet: bool = True):
            if not items:
                return
            doc.add_heading(heading, level=2)
            if bullet:
                for item in items:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                for item in items:
                    doc.add_paragraph(str(item))

        # Summary
        if "summary" in content:
            doc.add_heading("Резюме", level=2)
            doc.add_paragraph(content["summary"])

        # Risks
        risks = content.get("risks", [])
        if risks:
            doc.add_heading("Риски", level=2)
            for risk in risks:
                if isinstance(risk, dict):
                    p = doc.add_paragraph(style="List Bullet")
                    sev = risk.get("severity", "")
                    p.add_run(f"[{sev.upper()}] ").bold = True
                    p.add_run(risk.get("description", ""))

        add_section("Отсутствующие требования", content.get("missing_requirements", []))
        add_section("Вопросы заказчику", content.get("questions_to_client", []))
        add_section("Критерии приёмки", content.get("acceptance_criteria", []))
        add_section("Архитектурные риски", content.get("architecture_risks", []))

        # Эпик A4: диаграммы картинками
        if diagrams:
            doc.add_heading("Диаграммы", level=2)
            for d in diagrams:
                label = DIAGRAM_TYPE_LABELS.get(d.diagram_type, d.diagram_type)
                doc.add_heading(label, level=3)
                render_png = getattr(d, "render_png", None)
                render_status = getattr(d, "render_status", "pending")
                if render_png:
                    doc.add_picture(io.BytesIO(render_png), width=Inches(6))
                    if getattr(d, "standard_profile", None):
                        cap = doc.add_paragraph()
                        cap.add_run(
                            f"Стандарт: {STANDARD_LABELS.get(d.standard_profile, d.standard_profile)}"
                        ).italic = True
                else:
                    warn = doc.add_paragraph()
                    reason = {
                        "failed": "локальный рендер не удался",
                        "external_fallback": "рендер доступен только во внешнем сервисе — не встроен в документ",
                        "blocked_external": "рендер заблокирован политикой локального контура (ENFORCE_LOCAL_ONLY)",
                        "pending": "диаграмма ещё не отрендерена",
                    }.get(render_status, render_status)
                    warn.add_run(f"[Рендер недоступен: {reason}] Исходный код диаграммы:").bold = True
                    code_p = doc.add_paragraph(d.source_code)
                    code_p.style = doc.styles["Normal"]

        # Metadata
        conf = content.get("confidence", "")
        needs = content.get("needs_review", False)
        doc.add_heading("Метаданные", level=2)
        doc.add_paragraph(f"Уверенность: {conf}")
        doc.add_paragraph(f"Требует проверки: {'Да ⚠️' if needs else 'Нет'}")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # Fallback to plain text if python-docx not available
        text = f"{title}\n{'='*len(title)}\n\n"
        text += json.dumps(content, ensure_ascii=False, indent=2)
        return text.encode("utf-8")


def export_business_case_pdf(title: str, report) -> bytes:
    """Export business case to PDF using fpdf2 with Unicode support."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()

    _FONT_CANDIDATES = [
        r"C:\Windows\Fonts\DejaVuSans.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    _FONT_BOLD_CANDIDATES = [
        r"C:\Windows\Fonts\DejaVuSans-Bold.ttf",
        r"C:\Windows\Fonts\arialbd.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    ]
    font_path = next((f for f in _FONT_CANDIDATES if os.path.isfile(f)), None)
    font_bold_path = next((f for f in _FONT_BOLD_CANDIDATES if os.path.isfile(f)), None)

    if font_path:
        pdf.add_font("Custom", "", font_path, uni=True)
        if font_bold_path:
            pdf.add_font("Custom", "B", font_bold_path, uni=True)
        FNAME = "Custom"
    else:
        FNAME = "Helvetica"

    pdf.set_font(FNAME, "B", 16)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    e = report.latest_economic_estimate
    te = report.latest_task_estimate

    if e:
        pdf.set_font(FNAME, "B", 12)
        pdf.cell(0, 10, "Экономические показатели", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(FNAME, "", 10)

        items = [
            ("CAPEX", f"{e.capex:,.0f} руб."),
            ("OPEX/мес", f"{e.opex_monthly:,.0f} руб."),
            ("Выгода/мес", f"{e.benefit_monthly:,.0f} руб."),
            ("Окупаемость", f"{e.payback_months} мес." if e.payback_months > 0 else "Не окупается"),
            ("ROI 12 мес", f"{e.roi_12m_pct}%"),
        ]
        for label, val in items:
            pdf.set_font(FNAME, "B", 10)
            pdf.cell(60, 8, label, new_x="END")
            pdf.set_font(FNAME, "", 10)
            pdf.cell(0, 8, val, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    if te:
        pdf.set_font(FNAME, "B", 12)
        pdf.cell(0, 10, "Декомпозиция задач", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(FNAME, "", 10)
        pdf.cell(0, 8, f"Уверенность: {te.confidence}", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 8, f"Требует проверки: {'Да' if te.needs_review else 'Нет'}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        try:
            tj = json.loads(te.tasks_json) if isinstance(te.tasks_json, str) else te.tasks_json
            for t in tj.get("tasks", [])[:20]:
                pdf.set_font(FNAME, "", 9)
                line = f"  - {t.get('name','')[:60]} [{t.get('role','')}] {t.get('estimated_hours',0)}ч"
                pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
        except Exception:
            pass

    pdf.set_font(FNAME, "", 8)
    pdf.ln(8)
    pdf.cell(0, 6, f"Сгенерировано Analyst-Architect-AI {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT")

    return pdf.output()
